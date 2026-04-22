import asyncio
import io

from celery.utils.log import get_task_logger

from src.worker.app import app
from src.utils.logging import bind_job_context

logger = get_task_logger(__name__)


def _run_async(coro):
    loop = asyncio.new_event_loop()
    try:
        asyncio.set_event_loop(loop)
        return loop.run_until_complete(coro)
    finally:
        loop.close()


@app.task(
    bind=True,
    name="src.worker.tasks.summary.summary_task",
    queue="summary",
    max_retries=3,
)
def summary_task(self, transcription_id: str, user_id: int) -> dict:
    bind_job_context(transcription_id=transcription_id, user_id=user_id, kind="summary")
    return _run_async(_summary_async(transcription_id, user_id))


async def _summary_async(transcription_id: str, user_id: int) -> dict:
    from src.db.base import async_session_factory
    from src.db.repositories.transcription import get_transcription
    from src.services.summary import generate_summary
    from src.services.notification import send_message
    from src.bot.texts.ru import SUMMARY_READY, SUMMARY_ERROR

    async with async_session_factory() as session:
        transcription = await get_transcription(transcription_id, session)
        if not transcription or not transcription.result_text:
            await send_message(user_id, SUMMARY_ERROR)
            return {"status": "failed"}

        try:
            summary = await generate_summary(transcription.result_text)
            transcription.summary_text = summary
            await session.commit()
        except Exception as e:
            logger.error("summary_error user_id=%s err=%s", user_id, e, exc_info=True)
            await send_message(user_id, SUMMARY_ERROR)
            return {"status": "failed", "error": str(e)}

    await send_message(user_id, SUMMARY_READY.format(summary=summary), parse_mode="HTML")
    return {"status": "done"}


@app.task(
    bind=True,
    name="src.worker.tasks.summary.docx_task",
    queue="summary",
)
def docx_task(self, transcription_id: str, user_id: int) -> dict:
    bind_job_context(transcription_id=transcription_id, user_id=user_id, kind="docx")
    return _run_async(_docx_async(transcription_id, user_id))


async def _docx_async(transcription_id: str, user_id: int) -> dict:
    from src.db.base import async_session_factory
    from src.db.repositories.transcription import get_transcription
    from src.services.notification import send_document
    from docx import Document

    async with async_session_factory() as session:
        transcription = await get_transcription(transcription_id, session)
        if not transcription or not transcription.result_text:
            return {"status": "failed"}

    doc = Document()
    doc.add_heading("Транскрибация", 0)
    for para in transcription.result_text.split("\n"):
        doc.add_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    await send_document(user_id, buf.read(), "transcription.docx")
    return {"status": "done"}


@app.task(
    bind=True,
    name="src.worker.tasks.summary.srt_task",
    queue="summary",
)
def srt_task(self, transcription_id: str, user_id: int) -> dict:
    bind_job_context(transcription_id=transcription_id, user_id=user_id, kind="srt")
    return _run_async(_srt_async(transcription_id, user_id))


def _format_timestamp(seconds: float) -> str:
    """Format seconds as HH:MM:SS,mmm per SRT spec."""
    if seconds < 0:
        seconds = 0
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds - int(seconds)) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"


def _text_to_srt(text: str, total_duration: float) -> str:
    """Fallback SRT generator when segment-level timestamps are unavailable.

    Splits the text into roughly equal-duration sentences and emits an SRT
    file. Not as accurate as Groq's segment output but better than nothing
    for plain-text transcripts.
    """
    import re
    # Sentence split on ./!/? keeping non-empty fragments.
    parts = [p.strip() for p in re.split(r"(?<=[\.!?])\s+", text) if p.strip()]
    if not parts:
        parts = [text]
    if not total_duration or total_duration <= 0:
        total_duration = max(len(parts) * 3.0, 1.0)
    per = total_duration / len(parts)

    lines = []
    for idx, fragment in enumerate(parts, start=1):
        start = (idx - 1) * per
        end = min(idx * per, total_duration)
        lines.append(str(idx))
        lines.append(f"{_format_timestamp(start)} --> {_format_timestamp(end)}")
        lines.append(fragment)
        lines.append("")
    return "\n".join(lines)


async def _srt_async(transcription_id: str, user_id: int) -> dict:
    from src.db.base import async_session_factory
    from src.db.repositories.transcription import get_transcription
    from src.services.notification import send_document, send_message

    async with async_session_factory() as session:
        transcription = await get_transcription(transcription_id, session)
        if not transcription or not transcription.result_text:
            await send_message(user_id, "❌ Текст для субтитров не найден.")
            return {"status": "failed"}

        duration = float(transcription.duration_seconds or 0)
        srt = _text_to_srt(transcription.result_text, duration)

    await send_document(user_id, srt.encode("utf-8"), "subtitles.srt")
    return {"status": "done"}
